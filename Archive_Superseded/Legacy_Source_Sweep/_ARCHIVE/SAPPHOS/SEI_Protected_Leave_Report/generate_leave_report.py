import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd

# Paths
docx_path = r'C:\Users\mespindola\Desktop\SEI_Protected_Leave_Report\SEI_Protected_Leave_Report.docx'
chart_folder = r'C:\Users\mespindola\Desktop\SEI_Protected_Leave_Report\charts'
os.makedirs(chart_folder, exist_ok=True)

# Create DOCX
doc = Document()
doc.add_heading('SEI Protected Leave Executive Report', 0)
doc.add_paragraph('Prepared for Executive Leadership')

# --- Example Scenario: PDL ---
doc.add_heading('1. Example: Pregnancy Disability Leave (PDL)', level=1)
doc.add_paragraph('Employee: Jane Doe')
doc.add_paragraph('Role: Mid-level Environmental Consultant')
doc.add_paragraph('Total Available Hours (Monthly): 160')
doc.add_paragraph('Leave Taken: PDL (8 weeks / 320 hours over 2 months)')
doc.add_paragraph('Billable Hours Worked: 100')
doc.add_paragraph('Other Leaves: PTO – 8 hours')

doc.add_heading('Chargeability Calculation', level=2)
doc.add_paragraph('Adjusted Chargeability (%) = Chargeable Hours / (Total Available Hours – Protected Leave Hours) × 100')

total_hours = 160*2
protected_leave = 320  # PDL
company_leave = 8
billable_hours = 100
adjusted_chargeability = billable_hours / (total_hours + protected_leave - (protected_leave + company_leave)) * 100
doc.add_paragraph(f'Adjusted Chargeability ≈ {adjusted_chargeability:.1f}%')

# --- Payroll Code Matrix ---
doc.add_heading('2. Payroll Code Matrix: PDL', level=1)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Code'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'KPI Denominator Inclusion'
hdr_cells[3].text = 'Category'
hdr_cells[4].text = 'Notes / Best Practice'

codes = [
    ['PDL','Pregnancy Disability Leave','Exclude','Protected Leave','CA law: up to 4 months; job-protected; do not count in KPI denominator.'],
    ['PDL-Partial','Partial PDL (Intermittent)','Exclude','Protected Leave','Track exact hours/days; adjust denominator for intermittent absences.'],
    ['PDL-Return','Transition Back to Work','Include','Billable/Normal','Hours after leave counts normally; may include training/gradual ramp-up.'],
    ['PDL-Overlap','PDL concurrent with FMLA/CFRA','Exclude','Protected Leave','Ensure no double-counting; both leaves excluded from KPI denominator.']
]
for row in codes:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# --- Protected Leave Matrix ---
doc.add_heading('3. Full Protected Leave Matrix', level=1)
leave_codes = [
    ['FMLA','Family Medical Leave Act','Exclude','Protected Leave','Federal job-protected leave.'],
    ['CFRA','California Family Rights Act','Exclude','Protected Leave','State job-protected leave.'],
    ['PDL','Pregnancy Disability Leave','Exclude','Protected Leave','CA law, job-protected.'],
    ['ADA/FEHA','Disability Accommodation Leave','Exclude','Protected Leave','State/federal disability protections.'],
    ['Jury','Jury Duty','Exclude','Protected Leave','Statutory leave.'],
    ['Bereavement','Bereavement Leave','Exclude','Protected Leave','Company leave.'],
    ['Voting','Voting Leave','Exclude','Protected Leave','Statutory leave.'],
    ['Military','Military Leave','Exclude','Protected Leave','Statutory leave.'],
    ['Sick-Protected','Sick Leave – Job Protected','Exclude','Protected Leave','Sick leave covered under ADA/FMLA.'],
    ['Sick-NonProtected','Sick Leave – Non-Protected','Include','Non-Protected Leave','Company policy sick PTO.']
]
table2 = doc.add_table(rows=1, cols=5)
for i, col in enumerate(['Code','Description','KPI Denominator','Category','Notes']):
    table2.rows[0].cells[i].text = col
for row in leave_codes:
    cells = table2.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# --- Charts ---
# Bar Chart: Billable vs Non-Billable
plt.figure(figsize=(6,4))
sns.barplot(x=['Billable','Protected Leave','Company Leave'], y=[billable_hours, protected_leave, company_leave], palette='Blues')
plt.ylabel('Hours')
plt.title('Billable vs Non-Billable Distribution')
bar_chart_path = os.path.join(chart_folder,'billable_distribution.png')
plt.savefig(bar_chart_path)
plt.close()
doc.add_heading('4. Billable vs Non-Billable', level=1)
doc.add_picture(bar_chart_path, width=Inches(6))

# Heatmap: Leave Type Inclusion/Exclusion
leave_df = pd.DataFrame({
    'Leave Type':['FMLA','CFRA','PDL','ADA/FEHA','Jury','Bereavement','Voting','Military','Sick-Protected','Sick-NonProtected'],
    'KPI Included':[0,0,0,0,0,0,0,0,0,1]
})
plt.figure(figsize=(8,4))
sns.heatmap(leave_df[['KPI Included']], annot=True, yticklabels=leave_df['Leave Type'], cmap='coolwarm', cbar=False)
plt.title('Leave Type Inclusion (1=Included, 0=Excluded)')
heatmap_path = os.path.join(chart_folder,'leave_heatmap.png')
plt.savefig(heatmap_path)
plt.close()
doc.add_heading('5. Leave Type Inclusion Heatmap', level=1)
doc.add_picture(heatmap_path, width=Inches(6))

# Radar Chart: KPI coverage across categories
import numpy as np
categories = ['Chargeability','Work Quality','Client Satisfaction','Project Delivery','Professional Dev','Compliance & Teamwork']
values = [adjusted_chargeability, 25,25,20,15,10]
values += values[:1]  # close the loop
angles = np.linspace(0, 2*np.pi, len(categories)+1, endpoint=True)

plt.figure(figsize=(6,6))
ax = plt.subplot(111, polar=True)
ax.plot(angles, values, 'o-', linewidth=2)
ax.fill(angles, values, alpha=0.25)
ax.set_thetagrids(angles[:-1]*180/np.pi, categories)
plt.title('KPI Coverage Radar Chart')
radar_path = os.path.join(chart_folder,'kpi_radar.png')
plt.savefig(radar_path)
plt.close()
doc.add_heading('6. KPI Coverage Radar Chart', level=1)
doc.add_picture(radar_path, width=Inches(6))

# Save DOCX
doc.save(docx_path)
print(f"✅ DOCX report with charts generated at: {docx_path}")
