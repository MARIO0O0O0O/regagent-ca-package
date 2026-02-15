import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

# Paths
report_path = r'C:\Users\mespindola\Desktop\SEI_KPI_Executive_Report\SEI_KPI_HighLevel_Report.docx'
chart_folder = r'C:\Users\mespindola\Desktop\SEI_KPI_Executive_Report\charts'

# Create Document
doc = Document()
doc.add_heading('SEI Executive KPI Report', 0)

# Section 1: Overview
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'This report provides a high-level overview of SEI\'s key performance indicators (KPIs), '
    'the chargeability formula, protected and company leaves, and a workflow for managing performance evaluations. '
    'This report is intended for executive leadership and focuses on strategic insights rather than operational details.'
)

# Section 2: KPI Definitions
doc.add_heading('2. KPI Definitions', level=1)
doc.add_paragraph('KPI: Key Performance Indicator, a metric used to measure employee performance against business objectives.')
doc.add_paragraph('Chargeability: % of hours an employee works on client-billable tasks.')

doc.add_paragraph('Formula (Leave-Adjusted Chargeability):')
doc.add_paragraph('Adjusted Chargeability (%) = Chargeable Hours ÷ (Total Available Hours – Protected Leave Hours) × 100')

# Section 3: Leave Categories
doc.add_heading('3. Leave Categories', level=1)
doc.add_paragraph('Protected Leaves (cannot penalize):')
doc.add_paragraph('- FMLA (Family Medical Leave Act)')
doc.add_paragraph('- CFRA (California Family Rights Act)')
doc.add_paragraph('- Pregnancy Disability Leave (PDL)')
doc.add_paragraph('- ADA/FEHA approved accommodations')
doc.add_paragraph('- Paid Family Leave, Jury duty, Bereavement, Voting leave')

doc.add_paragraph('Company-Granted Leaves (also excluded from KPI calculations):')
doc.add_paragraph('- Vacation / PTO')
doc.add_paragraph('- Company Holidays')
doc.add_paragraph('- Training / Professional Development')

doc.add_paragraph('Included in KPI calculations:')
doc.add_paragraph('- Client-billable hours')
doc.add_paragraph('- Non-protected overtime')

# Section 4: High-Level Workflow
doc.add_heading('4. KPI Evaluation Workflow', level=1)
doc.add_paragraph(
    '1. Employees submit timesheets (billable and non-billable) weekly.\n'
    '2. HR collects leave data and flags protected leaves.\n'
    '3. Chargeability is calculated using the leave-adjusted formula.\n'
    '4. KPI tables are generated for executive review.\n'
    '5. Quarterly and annual summaries are presented to leadership.'
)

# Section 5: Sample Charts
doc.add_heading('5. Visual Insights', level=1)

# Generate KPI Bar Chart
kpi_chart_path = os.path.join(chart_folder, 'kpi_bar_chart.png')
employees = ['Alice', 'Bob', 'Charlie', 'David']
chargeability = [85, 78, 90, 70]

plt.figure(figsize=(6,4))
plt.bar(employees, chargeability, color='skyblue')
plt.title('High-Level Chargeability by Employee')
plt.ylabel('Chargeability (%)')
plt.ylim(0,100)
plt.savefig(kpi_chart_path)
plt.close()
doc.add_picture(kpi_chart_path, width=Inches(6))

# Generate Gantt-like Timeline (Simplified)
gantt_chart_path = os.path.join(chart_folder, 'gantt_timeline.png')
tasks = ['Policy Design','Chart Generation','Manager Training','Employee Communication','Quarterly Audit']
start = [0,2,4,6,8]
dur = [2,2,2,2,2]

plt.figure(figsize=(6,3))
for i, task in enumerate(tasks):
    plt.barh(task, dur[i], left=start[i], color='teal')
plt.title('High-Level KPI Implementation Timeline')
plt.xlabel('Months')
plt.savefig(gantt_chart_path)
plt.close()
doc.add_picture(gantt_chart_path, width=Inches(6))

# Section 6: Summary
doc.add_heading('6. Executive Summary', level=1)
doc.add_paragraph(
    '✅ KPI formula ensures fairness by adjusting for legally protected leaves.\n'
    '✅ Company leaves are also excluded to reflect true productivity.\n'
    '✅ High-level charts allow leadership to visualize chargeability trends and implementation timelines.\n'
    '✅ This report is suitable for strategic decision-making without going into operational details.'
)

# Save report
doc.save(report_path)
print(f"✅ DOCX report generated at: {report_path}")
