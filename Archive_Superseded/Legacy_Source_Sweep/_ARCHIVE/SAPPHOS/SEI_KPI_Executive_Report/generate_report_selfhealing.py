import sys
import subprocess

# --- Self-healing: install missing modules ---
modules = ['matplotlib', 'python-docx']
for m in modules:
    try:
        __import__(m)
    except ImportError:
        print(f'⚠ Module {m} missing. Installing locally...')
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', '--user', m])

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import os

# --- Paths ---
report_path = r'C:\Users\mespindola\Desktop\SEI_KPI_Executive_Report\SEI_KPI_HighLevel_Report.docx'
chart_folder = r'C:\Users\mespindola\Desktop\SEI_KPI_Executive_Report\charts'

# --- Create Document ---
doc = Document()
doc.add_heading('SEI Executive KPI Report', 0)

# Section 1: Overview
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'This high-level report provides executive leadership with insights into SEI\'s KPIs, '
    'chargeability formula, protected and company leaves, and evaluation workflow.'
)

# Section 2: KPI Definitions
doc.add_heading('2. KPI Definitions', level=1)
doc.add_paragraph('KPI: Key Performance Indicator, a metric used to measure employee performance against objectives.')
doc.add_paragraph('Chargeability: % of hours an employee works on client-billable tasks.')
doc.add_paragraph('Formula (Leave-Adjusted Chargeability):')
doc.add_paragraph('Adjusted Chargeability (%) = Chargeable Hours ÷ (Total Available Hours – Protected Leave Hours) × 100')

# Section 3: Leave Categories
doc.add_heading('3. Leave Categories', level=1)
doc.add_paragraph('Protected Leaves (cannot penalize):')
for leave in ['FMLA', 'CFRA', 'PDL', 'ADA/FEHA accommodations', 'Paid Family Leave', 'Jury duty', 'Bereavement', 'Voting leave']:
    doc.add_paragraph(f'- {leave}')

doc.add_paragraph('Company-Granted Leaves (excluded from KPI calculations):')
for leave in ['Vacation / PTO', 'Company Holidays', 'Training / Professional Development']:
    doc.add_paragraph(f'- {leave}')

doc.add_paragraph('Included in KPI calculations:')
for item in ['Client-billable hours', 'Non-protected overtime']:
    doc.add_paragraph(f'- {item}')

# Section 4: Workflow
doc.add_heading('4. KPI Evaluation Workflow', level=1)
workflow = [
    'Employees submit timesheets (billable and non-billable).',
    'HR collects leave data and flags protected leaves.',
    'Chargeability is calculated using leave-adjusted formula.',
    'KPI tables generated for executive review.',
    'Quarterly/annual summaries presented to leadership.'
]
for step in workflow:
    doc.add_paragraph(f'{step}')

# Section 5: Charts
doc.add_heading('5. Visual Insights', level=1)

# KPI Bar Chart
kpi_chart_path = os.path.join(chart_folder, 'kpi_bar_chart.png')
employees = ['Alice', 'Bob', 'Charlie', 'David']
chargeability = [85, 78, 90, 70]
plt.figure(figsize=(6,4))
plt.bar(employees, chargeability, color='skyblue')
plt.title('High-Level Chargeability by Employee')
plt.ylabel('Chargeability (%)')
plt.ylim(0,100)
plt.savefig(kpi_chart_path)
plt.close()
doc.add_picture(kpi_chart_path, width=Inches(6))

# Simplified Gantt Timeline
gantt_chart_path = os.path.join(chart_folder, 'gantt_timeline.png')
tasks = ['Policy Design','Chart Generation','Manager Training','Employee Communication','Quarterly Audit']
start = [0,2,4,6,8]
dur = [2,2,2,2,2]
plt.figure(figsize=(6,3))
for i, task in enumerate(tasks):
    plt.barh(task, dur[i], left=start[i], color='teal')
plt.title('KPI Implementation Timeline (Months)')
plt.xlabel('Months')
plt.savefig(gantt_chart_path)
plt.close()
doc.add_picture(gantt_chart_path, width=Inches(6))

# Section 6: Summary
doc.add_heading('6. Executive Summary', level=1)
doc.add_paragraph(
    '✅ Chargeability formula ensures fairness by adjusting for legally protected leaves.\n'
    '✅ Company leaves are excluded to reflect true productivity.\n'
    '✅ Visual aids provide leadership with clear trends and timelines.\n'
    '✅ Report is executive-focused and high-level.'
)

# Save DOCX
doc.save(report_path)
print(f"✅ DOCX report generated at: {report_path}")
