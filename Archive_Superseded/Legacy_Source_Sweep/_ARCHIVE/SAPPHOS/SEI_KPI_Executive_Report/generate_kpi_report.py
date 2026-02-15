import os
import sys

# Self-healing imports
try:
    import pandas as pd
except ImportError:
    print('⚠️ pandas not installed. CSV reading disabled.')
    pd = None

try:
    import matplotlib.pyplot as plt
except ImportError:
    print('⚠️ matplotlib not installed. Charts disabled.')
    plt = None

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    print('⚠️ python-docx not installed. DOCX generation disabled.')
    Document = None

# Paths
report_path = r'C:\Users\mespindola\Desktop\SEI_KPI_Executive_Report\SEI_KPI_HighLevel_Report.docx'
chart_folder = r'C:\Users\mespindola\Desktop\SEI_KPI_Executive_Report\charts'
csv_path = r'C:\Users\mespindola\Desktop\SEI_KPI_Executive_Report\chargeability_data.csv'

# Load CSV if pandas is available
if pd:
    if os.path.exists(csv_path):
        df = pd.read_csv(csv_path)
    else:
        df = None
else:
    df = None

# Create document if possible
if Document:
    doc = Document()
    doc.add_heading('SEI Executive KPI Report', 0)

    # Section 1: Overview
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "High-level executive report covering KPIs, chargeability formulas, "
        "protected leaves, and workflow for evaluation."
    )

    # Section 2: KPI Definitions
    doc.add_heading('2. KPI Definitions', level=1)
    doc.add_paragraph("Chargeability (% of hours billed to client)")
    doc.add_paragraph("Leave-Adjusted Chargeability formula: Chargeable Hours / (Available Hours - Protected Leave Hours) * 100")

    # Section 3: Leave Categories
    doc.add_heading('3. Leave Categories', level=1)
    doc.add_paragraph("Protected Leaves (excluded from KPI denominator): FMLA, CFRA, PDL, ADA/FEHA, Jury duty, Voting leave, Bereavement, Military leave")
    doc.add_paragraph("Company Leaves (also excluded): Vacation/PTO, Company Holidays, Training")
    doc.add_paragraph("Non-Protected Leaves (included in KPI denominator): Discretionary unpaid leave, non-statutory sick leave, partial-day tardies")

    # Section 4: KPI Evaluation Workflow
    doc.add_heading('4. KPI Evaluation Workflow', level=1)
    doc.add_paragraph(
        "1. Employees submit timesheets weekly.\n"
        "2. HR collects leave data.\n"
        "3. Leave-adjusted chargeability is calculated.\n"
        "4. KPI tables generated for leadership.\n"
        "5. Quarterly and annual summaries presented to executives."
    )

    # Section 5: Industry Benchmarks & Broader KPIs
    doc.add_heading('5. Industry Benchmarks & Broader KPIs', level=1)
    doc.add_paragraph(
        "Technical/Junior: 75–85%\n"
        "Mid-Level/PM: 75–90%\n"
        "Leadership/Partner: 40–75%\n"
        "Firm averages: 65–80%\n\n"
        "Broader KPIs: Revenue per billable employee, Billable realization rate, Project margin, Client satisfaction."
    )

    # Section 6: Fairness & Equity
    doc.add_heading('6. Fairness & Equity', level=1)
    doc.add_paragraph(
        "Role-aligned expectations, tracking non-billable strategic work, clear definitions, quarterly audits, and compliance with protected leave legislation."
    )

    # Section 7: Charts
    if plt and df is not None:
        # KPI Bar Chart
        kpi_chart = os.path.join(chart_folder, 'kpi_bar_chart.png')
        plt.figure(figsize=(6,4))
        plt.bar(df['Employee'], df['ChargeableHours'] / (df['AvailableHours'] - df['ProtectedLeaveHours']) * 100, color='skyblue')
        plt.ylim(0,100)
        plt.ylabel('Chargeability (%)')
        plt.title('High-Level Chargeability by Employee')
        plt.savefig(kpi_chart)
        plt.close()
        doc.add_heading('7. Visual Insights', level=1)
        doc.add_picture(kpi_chart, width=Inches(6))

        # Gantt-like timeline
        gantt_chart = os.path.join(chart_folder, 'gantt_timeline.png')
        tasks = ['Policy Design','Chart Generation','Manager Training','Employee Communication','Quarterly Audit']
        start = [0,2,4,6,8]
        dur = [2,2,2,2,2]
        plt.figure(figsize=(6,3))
        for i, task in enumerate(tasks):
            plt.barh(task, dur[i], left=start[i], color='teal')
        plt.title('KPI Implementation Timeline')
        plt.xlabel('Months')
        plt.savefig(gantt_chart)
        plt.close()
        doc.add_picture(gantt_chart, width=Inches(6))

    # Section 8: Executive Summary
    doc.add_heading('8. Executive Summary', level=1)
    doc.add_paragraph(
        "✅ KPI formula adjusts for protected and company leaves.\n"
        "✅ Role-specific targets increase transparency.\n"
        "✅ Visuals allow executive quick-read insights.\n"
        "✅ Ensures fairness and compliance."
    )

    # Save document
    doc.save(report_path)
    print(f"✅ DOCX report generated at: {report_path}")
else:
    print("⚠️ DOCX generation skipped because python-docx is not installed.")
